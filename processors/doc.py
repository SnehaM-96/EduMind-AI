from docx import Document


def extract_docx_text(docx_file):

    try:

        doc = Document(docx_file)
        text = ""

        for para in doc.paragraphs:

            if para.text.strip():
                text += para.text + "\n"

        # Also extract text from tables
        for table in doc.tables:

            for row in table.rows:

                row_text = " | ".join(
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                )

                if row_text:
                    text += row_text + "\n"

        return text.strip()

    except Exception :
        return ""